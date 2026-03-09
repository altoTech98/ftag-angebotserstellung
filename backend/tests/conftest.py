"""
Shared fixtures for Frank Türen AG test suite.
"""

import os
import sys
import io
import pytest
from unittest.mock import MagicMock, patch

# Add backend to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

# Set testing environment before importing anything
os.environ["TESTING"] = "true"
os.environ["ENVIRONMENT"] = "development"


@pytest.fixture
def sample_position():
    """A typical door position from a tender document."""
    return {
        "position": "1.01",
        "beschreibung": "Stahltür einflügelig mit Brandschutz",
        "tuertyp": "Stahltür",
        "brandschutz": "EI30",
        "einbruchschutz": "",
        "schallschutz": "",
        "breite": 900,
        "hoehe": 2100,
        "menge": 2,
        "einheit": "Stk",
    }


@pytest.fixture
def sample_positions():
    """Multiple door positions for batch testing."""
    return [
        {
            "position": "1.01",
            "beschreibung": "Stahltür mit Brandschutz T30",
            "tuertyp": "Stahltür",
            "brandschutz": "T30",
            "einbruchschutz": "",
            "menge": 3,
            "einheit": "Stk",
            "breite": 900,
            "hoehe": 2100,
        },
        {
            "position": "1.02",
            "beschreibung": "Holztür Innentür",
            "tuertyp": "Holztür",
            "brandschutz": "",
            "einbruchschutz": "",
            "menge": 5,
            "einheit": "Stk",
            "breite": 800,
            "hoehe": 2000,
        },
        {
            "position": "1.03",
            "beschreibung": "Sicherheitstür RC3 mit Brandschutz EI60",
            "tuertyp": "Stahltür",
            "brandschutz": "EI60",
            "einbruchschutz": "RC3",
            "menge": 1,
            "einheit": "Stk",
            "breite": 1100,
            "hoehe": 2200,
        },
    ]


@pytest.fixture
def sample_requirements(sample_positions):
    """A full requirements dict as returned by the analyze pipeline."""
    return {
        "projekt": "Testprojekt Schulhaus Buochs",
        "auftraggeber": "Gemeinde Buochs",
        "positionen": sample_positions,
    }


@pytest.fixture
def sample_matched_positions(sample_positions):
    """Matched positions as returned by the product matcher."""
    return [
        {
            "status": "matched",
            "confidence": 0.85,
            "position": "1.01",
            "beschreibung": "Stahltür mit Brandschutz T30",
            "menge": 3,
            "einheit": "Stk",
            "matched_products": [{"bezeichnung": "FTAG Stahltür T30"}],
            "reason": "Produkt im FTAG-Sortiment verfügbar",
            "original_position": sample_positions[0],
        },
        {
            "status": "matched",
            "confidence": 0.90,
            "position": "1.02",
            "beschreibung": "Holztür Innentür",
            "menge": 5,
            "einheit": "Stk",
            "matched_products": [{"bezeichnung": "FTAG Holztür Standard"}],
            "reason": "Produkt im FTAG-Sortiment verfügbar",
            "original_position": sample_positions[1],
        },
    ]


@pytest.fixture
def sample_unmatched_positions(sample_positions):
    """Unmatched positions for gap report testing."""
    return [
        {
            "status": "unmatched",
            "confidence": 0.0,
            "position": "1.03",
            "beschreibung": "Sicherheitstür RC3 mit Brandschutz EI60",
            "menge": 1,
            "einheit": "Stk",
            "matched_products": [],
            "reason": "Kein exaktes Produkt gefunden",
            "original_position": sample_positions[2],
        },
    ]


@pytest.fixture
def sample_txt_bytes():
    """Simple text file content for document parser tests."""
    return "Ausschreibung Türen\nPosition 1: Stahltür T30\nPosition 2: Holztür innen".encode("utf-8")


@pytest.fixture
def sample_excel_bytes():
    """Minimal Excel file bytes for parser tests."""
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Türliste"
    ws.append(["Position", "Beschreibung", "Menge", "Einheit"])
    ws.append(["1.01", "Stahltür T30", 3, "Stk"])
    ws.append(["1.02", "Holztür innen", 5, "Stk"])

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


@pytest.fixture
def sample_word_bytes():
    """Minimal Word .docx file bytes for parser tests."""
    from docx import Document

    doc = Document()
    doc.add_paragraph("Ausschreibung Türen")
    doc.add_paragraph("Position 1: Stahltür mit Brandschutz T30")
    doc.add_paragraph("Position 2: Holztür innen")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
