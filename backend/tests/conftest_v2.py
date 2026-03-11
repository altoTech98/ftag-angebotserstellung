"""
Shared fixtures for v2 tests.

Provides sample document bytes, mock ParseResult, and sample schema instances
for use across all v2 test modules.
"""

import io

import pytest

from v2.parsers.base import ParseResult
from v2.schemas.common import (
    BrandschutzKlasse,
    FieldSource,
    MaterialTyp,
    OeffnungsArt,
    SchallschutzKlasse,
    ZargenTyp,
)
from v2.schemas.extraction import ExtractedDoorPosition
from v2.schemas.matching import (
    DimensionScore,
    MatchCandidate,
    MatchDimension,
    MatchResult,
)
from v2.schemas.adversarial import (
    AdversarialCandidate,
    AdversarialResult,
    DimensionCoT,
    ValidationStatus,
)
from v2.schemas.gaps import (
    AlternativeProduct,
    GapDimension,
    GapItem,
    GapReport,
    GapSeverity,
)


# ---------------------------------------------------------------------------
# Phase 3: Cross-document intelligence fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def xlsx_positions() -> list[ExtractedDoorPosition]:
    """3 door positions from an XLSX door list.

    Positions 1.01, 1.02, 1.03 with basic dims/material but no fire rating.
    """
    return [
        ExtractedDoorPosition(
            positions_nr="1.01",
            positions_bezeichnung="Buerotuere",
            raum_nr="B101",
            geschoss="EG",
            breite_mm=1000,
            hoehe_mm=2100,
            material_blatt=MaterialTyp.HOLZ,
            quellen={
                "breite_mm": FieldSource(dokument="tuerliste.xlsx", zeile=2, zelle="D2", sheet="Tuerliste", konfidenz=0.95),
                "hoehe_mm": FieldSource(dokument="tuerliste.xlsx", zeile=2, zelle="E2", sheet="Tuerliste", konfidenz=0.95),
                "material_blatt": FieldSource(dokument="tuerliste.xlsx", zeile=2, zelle="H2", sheet="Tuerliste", konfidenz=0.9),
            },
        ),
        ExtractedDoorPosition(
            positions_nr="1.02",
            positions_bezeichnung="Flurtuere",
            raum_nr="F201",
            geschoss="OG",
            breite_mm=900,
            hoehe_mm=2050,
            material_blatt=MaterialTyp.STAHL,
            quellen={
                "breite_mm": FieldSource(dokument="tuerliste.xlsx", zeile=3, zelle="D3", sheet="Tuerliste", konfidenz=0.95),
                "hoehe_mm": FieldSource(dokument="tuerliste.xlsx", zeile=3, zelle="E3", sheet="Tuerliste", konfidenz=0.95),
            },
        ),
        ExtractedDoorPosition(
            positions_nr="1.03",
            positions_bezeichnung="WC-Tuere",
            raum_nr="WC01",
            geschoss="EG",
            breite_mm=800,
            hoehe_mm=2000,
            material_blatt=MaterialTyp.HOLZ,
            quellen={
                "breite_mm": FieldSource(dokument="tuerliste.xlsx", zeile=4, zelle="D4", sheet="Tuerliste", konfidenz=0.95),
            },
        ),
    ]


@pytest.fixture
def pdf_positions() -> list[ExtractedDoorPosition]:
    """2 door positions from a PDF spec document.

    Position 'Pos. 1.01' with fire rating (overlaps xlsx 1.01).
    Position 'Tuer 1.02' with detailed specs (overlaps xlsx 1.02).
    Different naming convention to test normalization.
    """
    return [
        ExtractedDoorPosition(
            positions_nr="Pos. 1.01",
            brandschutz_klasse=BrandschutzKlasse.EI30,
            rauchschutz=True,
            schallschutz_klasse=SchallschutzKlasse.RW_32,
            schallschutz_db=32,
            quellen={
                "brandschutz_klasse": FieldSource(dokument="spec.pdf", seite=3, konfidenz=0.95),
                "rauchschutz": FieldSource(dokument="spec.pdf", seite=3, konfidenz=0.9),
                "schallschutz_klasse": FieldSource(dokument="spec.pdf", seite=4, konfidenz=0.85),
            },
        ),
        ExtractedDoorPosition(
            positions_nr="Tuer 1.02",
            brandschutz_klasse=BrandschutzKlasse.EI60,
            wandstaerke_mm=180,
            drueckergarnitur="Edelstahl gebuerstet",
            quellen={
                "brandschutz_klasse": FieldSource(dokument="spec.pdf", seite=5, konfidenz=0.95),
                "wandstaerke_mm": FieldSource(dokument="spec.pdf", seite=5, konfidenz=0.9),
            },
        ),
    ]


@pytest.fixture
def conflicting_positions() -> list[ExtractedDoorPosition]:
    """2 positions with conflicting brandschutz_klasse (T30 vs T90)."""
    return [
        ExtractedDoorPosition(
            positions_nr="1.01",
            brandschutz_klasse=BrandschutzKlasse.T30,
            breite_mm=1000,
            hoehe_mm=2100,
            quellen={
                "brandschutz_klasse": FieldSource(dokument="a.xlsx", konfidenz=0.9),
                "breite_mm": FieldSource(dokument="a.xlsx", konfidenz=0.95),
            },
        ),
        ExtractedDoorPosition(
            positions_nr="1.01",
            brandschutz_klasse=BrandschutzKlasse.T90,
            breite_mm=1000,
            hoehe_mm=2100,
            quellen={
                "brandschutz_klasse": FieldSource(dokument="b.pdf", konfidenz=0.95),
                "breite_mm": FieldSource(dokument="b.pdf", konfidenz=0.9),
            },
        ),
    ]


@pytest.fixture
def general_spec_text() -> str:
    """Sample general specification text from a DOCX Pflichtenheft."""
    return "Alle Innentüren im OG müssen mindestens T30 Brandschutz aufweisen"


# ---------------------------------------------------------------------------
# Phase 1-2 fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    """Minimal valid PDF bytes for testing.

    Creates a tiny 1-page PDF with a simple text content.
    This is a hardcoded minimal valid PDF structure.
    """
    # Minimal valid PDF 1.0 structure
    pdf_content = (
        b"%PDF-1.0\n"
        b"1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        b"2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        b"3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>\nendobj\n"
        b"4 0 obj\n<< /Length 44 >>\nstream\n"
        b"BT /F1 12 Tf 100 700 Td (Tuer Nr. 1.01) Tj ET\n"
        b"endstream\nendobj\n"
        b"5 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
        b"xref\n0 6\n"
        b"0000000000 65535 f \n"
        b"0000000009 00000 n \n"
        b"0000000058 00000 n \n"
        b"0000000115 00000 n \n"
        b"0000000266 00000 n \n"
        b"0000000360 00000 n \n"
        b"trailer\n<< /Size 6 /Root 1 0 R >>\n"
        b"startxref\n434\n%%EOF"
    )
    return pdf_content


@pytest.fixture
def sample_docx_bytes() -> bytes:
    """Minimal valid DOCX bytes for testing.

    Uses python-docx to create in-memory DOCX with a paragraph and table.
    """
    from docx import Document

    doc = Document()
    doc.add_paragraph("Tuerliste Projekt Test")

    table = doc.add_table(rows=3, cols=4)
    headers = ["Tuer Nr.", "Breite", "Hoehe", "Brandschutz"]
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    # Row 1
    for i, val in enumerate(["1.01", "1000", "2100", "EI30"]):
        table.rows[1].cells[i].text = val
    # Row 2
    for i, val in enumerate(["1.02", "900", "2050", "EI60"]):
        table.rows[2].cells[i].text = val

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def sample_xlsx_bytes() -> bytes:
    """Minimal valid XLSX bytes for testing.

    Creates an in-memory workbook simulating a Tuerliste
    with columns like Tuer Nr., Breite, Hoehe, Brandschutz.
    """
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = "Tuerliste"

    # Header row
    headers = [
        "Tuer Nr.", "Bezeichnung", "Raum", "Breite (mm)",
        "Hoehe (mm)", "Brandschutz", "Schallschutz (dB)",
        "Material", "Anzahl"
    ]
    ws.append(headers)

    # Data rows
    ws.append(["1.01", "Buerotuere", "B101", 1000, 2100, "EI30", 32, "Holz", 1])
    ws.append(["1.02", "Flurtuere", "F201", 900, 2050, "EI60", 37, "Stahl", 2])
    ws.append(["1.03", "WC-Tuere", "WC01", 800, 2000, "", 27, "Holz", 1])

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


@pytest.fixture
def mock_parse_result() -> ParseResult:
    """Returns a ParseResult with sample data."""
    return ParseResult(
        text="| Tuer Nr. | Breite | Hoehe | Brandschutz |\n|---|---|---|---|\n| 1.01 | 1000 | 2100 | EI30 |",
        format="xlsx",
        page_count=1,
        warnings=[],
        metadata={"sheet_names": ["Tuerliste"], "method": "openpyxl"},
        source_file="test_tuerliste.xlsx",
        tables=["| Tuer Nr. | Breite | Hoehe | Brandschutz |\n|---|---|---|---|\n| 1.01 | 1000 | 2100 | EI30 |"],
    )


@pytest.fixture
def sample_door_position() -> ExtractedDoorPosition:
    """Returns an ExtractedDoorPosition with realistic test data."""
    return ExtractedDoorPosition(
        positions_nr="1.01",
        positions_bezeichnung="Buerotuere Typ A",
        raum_nr="B101",
        raum_bezeichnung="Buero Geschaeftsleitung",
        geschoss="EG",
        breite_mm=1000,
        hoehe_mm=2100,
        wandstaerke_mm=150,
        brandschutz_klasse=BrandschutzKlasse.EI30,
        rauchschutz=True,
        schallschutz_klasse=SchallschutzKlasse.RW_32,
        schallschutz_db=32,
        material_blatt=MaterialTyp.HOLZ,
        material_zarge=ZargenTyp.UMFASSUNGSZARGE,
        oeffnungsart=OeffnungsArt.DREHFLUEGEL,
        anschlag_richtung="DIN links",
        oberflaeche="CPL Weiss",
        drueckergarnitur="Edelstahl gebuerstet",
        schlossart="Einsteckschloss BB",
        anzahl=1,
        quellen={
            "breite_mm": FieldSource(
                dokument="tuerliste.xlsx",
                zeile=5,
                zelle="D5",
                sheet="Tuerliste",
                konfidenz=0.95,
            ),
            "brandschutz_klasse": FieldSource(
                dokument="tuerliste.xlsx",
                zeile=5,
                zelle="F5",
                sheet="Tuerliste",
                konfidenz=0.99,
            ),
        },
    )


# ---------------------------------------------------------------------------
# Phase 7: Excel output generation fixtures
# ---------------------------------------------------------------------------


@pytest.fixture
def sample_positions() -> list[ExtractedDoorPosition]:
    """3 sample positions for Excel output tests."""
    return [
        ExtractedDoorPosition(
            positions_nr="1.01",
            positions_bezeichnung="Buerotuere EG",
            quellen={
                "breite_mm": FieldSource(dokument="tuerliste.xlsx", zeile=2, konfidenz=0.95),
            },
        ),
        ExtractedDoorPosition(
            positions_nr="1.02",
            positions_bezeichnung="Flurtuere OG",
            quellen={
                "breite_mm": FieldSource(dokument="spec.pdf", seite=3, konfidenz=0.9),
            },
        ),
        ExtractedDoorPosition(
            positions_nr="1.03",
            positions_bezeichnung="Kellertuere UG",
            quellen={},
        ),
    ]


def _make_dimension_scores(scores: dict[str, float]) -> list[DimensionScore]:
    """Helper to create DimensionScore list from {dim_name: score} dict."""
    return [
        DimensionScore(
            dimension=MatchDimension(dim),
            score=score,
            begruendung=f"Bewertung fuer {dim}: {score:.0%} Uebereinstimmung",
        )
        for dim, score in scores.items()
    ]


@pytest.fixture
def sample_match_results() -> list[MatchResult]:
    """3 MatchResult objects: high-confidence, partial, no-match."""
    return [
        MatchResult(
            positions_nr="1.01",
            bester_match=MatchCandidate(
                produkt_id="P-100",
                produkt_name="Buerotuere Standard EI30",
                gesamt_konfidenz=0.95,
                dimension_scores=_make_dimension_scores({
                    "Masse": 0.98, "Brandschutz": 0.95, "Schallschutz": 0.90,
                    "Material": 0.97, "Zertifizierung": 0.92, "Leistung": 0.88,
                }),
                begruendung="Sehr gute Uebereinstimmung bei allen Dimensionen",
            ),
            alternative_matches=[],
            hat_match=True,
            match_methode="tfidf_ai",
        ),
        MatchResult(
            positions_nr="1.02",
            bester_match=MatchCandidate(
                produkt_id="P-200",
                produkt_name="Flurtuere Brandschutz EI60",
                gesamt_konfidenz=0.72,
                dimension_scores=_make_dimension_scores({
                    "Masse": 0.85, "Brandschutz": 0.60, "Schallschutz": 0.70,
                    "Material": 0.80, "Zertifizierung": 0.65, "Leistung": 0.72,
                }),
                begruendung="Teilweise Uebereinstimmung, Brandschutz abweichend",
            ),
            alternative_matches=[],
            hat_match=True,
            match_methode="tfidf_ai",
        ),
        MatchResult(
            positions_nr="1.03",
            bester_match=MatchCandidate(
                produkt_id="P-300",
                produkt_name="Kellertuere Basis",
                gesamt_konfidenz=0.35,
                dimension_scores=_make_dimension_scores({
                    "Masse": 0.40, "Brandschutz": 0.20, "Schallschutz": 0.30,
                    "Material": 0.50, "Zertifizierung": 0.25, "Leistung": 0.45,
                }),
                begruendung="Keine passende Uebereinstimmung",
            ),
            alternative_matches=[],
            hat_match=False,
            match_methode="tfidf_ai",
        ),
    ]


def _make_dimension_cot(scores: dict[str, tuple[float, str]]) -> list[DimensionCoT]:
    """Helper: {dim_name: (score, reasoning)} -> list[DimensionCoT]."""
    return [
        DimensionCoT(
            dimension=dim,
            score=score,
            reasoning=reasoning,
            confidence_level="hoch" if score > 0.9 else "niedrig",
        )
        for dim, (score, reasoning) in scores.items()
    ]


@pytest.fixture
def sample_adversarial_results() -> list[AdversarialResult]:
    """3 AdversarialResult objects: bestaetigt (0.97), unsicher (0.75), abgelehnt (0.40)."""
    return [
        AdversarialResult(
            positions_nr="1.01",
            validation_status=ValidationStatus.BESTAETIGT,
            adjusted_confidence=0.97,
            bester_match=AdversarialCandidate(
                produkt_id="P-100",
                produkt_name="Buerotuere Standard EI30",
                adjusted_confidence=0.97,
                dimension_scores=_make_dimension_cot({
                    "Masse": (0.98, "Breite und Hoehe stimmen exakt ueberein"),
                    "Brandschutz": (0.96, "EI30 Anforderung vollstaendig erfuellt"),
                    "Schallschutz": (0.92, "32 dB Anforderung erfuellt"),
                    "Material": (0.97, "Holz wie gefordert"),
                    "Zertifizierung": (0.95, "CE Kennzeichnung vorhanden"),
                    "Leistung": (0.90, "Alle Leistungsmerkmale erfuellt"),
                }),
                reasoning_summary="Produkt erfuellt alle Anforderungen vollstaendig",
            ),
            debate=[],
            resolution_reasoning="FOR und AGAINST stimmen ueberein: optimale Zuordnung",
            per_dimension_cot=_make_dimension_cot({
                "Masse": (0.98, "Breite und Hoehe stimmen exakt ueberein"),
                "Brandschutz": (0.96, "EI30 Anforderung vollstaendig erfuellt"),
                "Schallschutz": (0.92, "32 dB Anforderung erfuellt"),
                "Material": (0.97, "Holz wie gefordert"),
                "Zertifizierung": (0.95, "CE Kennzeichnung vorhanden"),
                "Leistung": (0.90, "Alle Leistungsmerkmale erfuellt"),
            }),
        ),
        AdversarialResult(
            positions_nr="1.02",
            validation_status=ValidationStatus.UNSICHER,
            adjusted_confidence=0.75,
            bester_match=AdversarialCandidate(
                produkt_id="P-200",
                produkt_name="Flurtuere Brandschutz EI60",
                adjusted_confidence=0.75,
                dimension_scores=_make_dimension_cot({
                    "Masse": (0.85, "Abmessungen leicht abweichend"),
                    "Brandschutz": (0.62, "EI60 gefordert, nur EI30 verfuegbar im naechsten Modell"),
                    "Schallschutz": (0.70, "Schallschutz knapp unter Anforderung"),
                    "Material": (0.80, "Material grundsaetzlich passend"),
                    "Zertifizierung": (0.68, "Zertifizierung nur teilweise vorhanden"),
                    "Leistung": (0.72, "Leistung teilweise erfuellt"),
                }),
                reasoning_summary="Teilweise Uebereinstimmung mit Abweichungen bei Brandschutz",
            ),
            debate=[],
            resolution_reasoning="Brandschutz-Abweichung fuehrt zu Unsicherheit",
            per_dimension_cot=_make_dimension_cot({
                "Masse": (0.85, "Abmessungen leicht abweichend"),
                "Brandschutz": (0.62, "EI60 gefordert, nur EI30 verfuegbar im naechsten Modell"),
                "Schallschutz": (0.70, "Schallschutz knapp unter Anforderung"),
                "Material": (0.80, "Material grundsaetzlich passend"),
                "Zertifizierung": (0.68, "Zertifizierung nur teilweise vorhanden"),
                "Leistung": (0.72, "Leistung teilweise erfuellt"),
            }),
        ),
        AdversarialResult(
            positions_nr="1.03",
            validation_status=ValidationStatus.ABGELEHNT,
            adjusted_confidence=0.40,
            bester_match=AdversarialCandidate(
                produkt_id="P-300",
                produkt_name="Kellertuere Basis",
                adjusted_confidence=0.40,
                dimension_scores=_make_dimension_cot({
                    "Masse": (0.40, "Abmessungen passen nicht"),
                    "Brandschutz": (0.20, "Kein Brandschutz vorhanden"),
                    "Schallschutz": (0.30, "Schallschutz nicht ausreichend"),
                    "Material": (0.50, "Material nur teilweise passend"),
                    "Zertifizierung": (0.25, "Fehlende Zertifizierungen"),
                    "Leistung": (0.45, "Leistung unzureichend"),
                }),
                reasoning_summary="Keine passende Zuordnung moeglich",
            ),
            debate=[],
            resolution_reasoning="Zu viele Abweichungen, kein brauchbarer Match",
            per_dimension_cot=_make_dimension_cot({
                "Masse": (0.40, "Abmessungen passen nicht"),
                "Brandschutz": (0.20, "Kein Brandschutz vorhanden"),
                "Schallschutz": (0.30, "Schallschutz nicht ausreichend"),
                "Material": (0.50, "Material nur teilweise passend"),
                "Zertifizierung": (0.25, "Fehlende Zertifizierungen"),
                "Leistung": (0.45, "Leistung unzureichend"),
            }),
        ),
    ]


@pytest.fixture
def sample_gap_reports() -> list[GapReport]:
    """GapReport for unsicher (1.02) and abgelehnt (1.03) positions only."""
    return [
        GapReport(
            positions_nr="1.02",
            gaps=[
                GapItem(
                    dimension=GapDimension.BRANDSCHUTZ,
                    schweregrad=GapSeverity.KRITISCH,
                    anforderung_wert="EI60",
                    katalog_wert="EI30",
                    abweichung_beschreibung="Brandschutzklasse nicht ausreichend",
                    kundenvorschlag="Upgrade auf EI60-Modell moeglich",
                    technischer_hinweis="Tuerblattwechsel erforderlich",
                    gap_geschlossen_durch=["P-201"],
                ),
                GapItem(
                    dimension=GapDimension.SCHALLSCHUTZ,
                    schweregrad=GapSeverity.MAJOR,
                    anforderung_wert="37 dB",
                    katalog_wert="32 dB",
                    abweichung_beschreibung="5 dB unter Anforderung",
                    kundenvorschlag="Schallschutz-Upgrade erhaeltlich",
                    technischer_hinweis=None,
                    gap_geschlossen_durch=[],
                ),
            ],
            alternativen=[
                AlternativeProduct(
                    produkt_id="P-201",
                    produkt_name="Flurtuere Premium EI60",
                    teilweise_deckung=0.85,
                    verbleibende_gaps=["Schallschutz"],
                    geschlossene_gaps=["Brandschutz"],
                ),
            ],
            zusammenfassung="Brandschutz und Schallschutz nicht vollstaendig erfuellt",
            validation_status="unsicher",
        ),
        GapReport(
            positions_nr="1.03",
            gaps=[
                GapItem(
                    dimension=GapDimension.BRANDSCHUTZ,
                    schweregrad=GapSeverity.KRITISCH,
                    anforderung_wert="EI90",
                    katalog_wert=None,
                    abweichung_beschreibung="Kein Brandschutz im Basismodell",
                    kundenvorschlag="Sonderanfertigung erforderlich",
                    technischer_hinweis="Komplett anderes Tuersystem noetig",
                    gap_geschlossen_durch=[],
                ),
                GapItem(
                    dimension=GapDimension.MASSE,
                    schweregrad=GapSeverity.MAJOR,
                    anforderung_wert="1200x2400mm",
                    katalog_wert="1000x2100mm",
                    abweichung_beschreibung="Uebergroesse nicht im Standardsortiment",
                    kundenvorschlag=None,
                    technischer_hinweis="Sondermass moeglich gegen Aufpreis",
                    gap_geschlossen_durch=["P-301"],
                ),
                GapItem(
                    dimension=GapDimension.MATERIAL,
                    schweregrad=GapSeverity.MINOR,
                    anforderung_wert="Edelstahl V2A",
                    katalog_wert="Stahl verzinkt",
                    abweichung_beschreibung="Material weicht ab",
                    kundenvorschlag=None,
                    technischer_hinweis=None,
                    gap_geschlossen_durch=[],
                ),
            ],
            alternativen=[
                AlternativeProduct(
                    produkt_id="P-301",
                    produkt_name="Kellertuere Sonder",
                    teilweise_deckung=0.55,
                    verbleibende_gaps=["Brandschutz", "Material"],
                    geschlossene_gaps=["Masse"],
                ),
            ],
            zusammenfassung="Mehrere kritische Abweichungen, Sonderanfertigung empfohlen",
            validation_status="abgelehnt",
        ),
    ]
