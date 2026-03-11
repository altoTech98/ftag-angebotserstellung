"""
Tests for v2 DOCX parser.

Tests DOCX parsing with paragraph extraction, heading markers, and table handling.
All tests verify the parser returns ParseResult and never raises.
"""

import io

import pytest
from docx import Document

from v2.parsers.base import ParseResult


def _make_docx_bytes(paragraphs=None, tables=None, headings=None) -> bytes:
    """Helper to create DOCX bytes with specified content."""
    doc = Document()

    if headings:
        for level, text in headings:
            doc.add_heading(text, level=level)

    if paragraphs:
        for text in paragraphs:
            doc.add_paragraph(text)

    if tables:
        for table_data in tables:
            rows = len(table_data)
            cols = len(table_data[0]) if table_data else 0
            table = doc.add_table(rows=rows, cols=cols)
            for r_idx, row in enumerate(table_data):
                for c_idx, val in enumerate(row):
                    table.rows[r_idx].cells[c_idx].text = str(val)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxParagraphExtraction:
    """Test basic paragraph extraction."""

    def test_docx_paragraph_extraction(self):
        """Parse valid DOCX -> text contains paragraph content."""
        from v2.parsers.docx_parser import parse_docx

        docx_bytes = _make_docx_bytes(
            paragraphs=["Tuerliste Projekt Alpha", "Spezifikation fuer 50 Tueren"]
        )
        result = parse_docx(docx_bytes, filename="test.docx")

        assert isinstance(result, ParseResult)
        assert "Tuerliste Projekt Alpha" in result.text
        assert "Spezifikation" in result.text
        assert result.format == "docx"

    def test_docx_returns_parse_result(self):
        """Return type is ParseResult."""
        from v2.parsers.docx_parser import parse_docx

        docx_bytes = _make_docx_bytes(paragraphs=["Test"])
        result = parse_docx(docx_bytes, filename="test.docx")
        assert isinstance(result, ParseResult)


class TestDocxTableExtraction:
    """Test table extraction from DOCX."""

    def test_docx_table_extraction(self):
        """Parse DOCX with table -> text contains table content, tables list is non-empty."""
        from v2.parsers.docx_parser import parse_docx

        table_data = [
            ["Tuer Nr.", "Breite", "Hoehe", "Brandschutz"],
            ["1.01", "1000", "2100", "EI30"],
            ["1.02", "900", "2050", "EI60"],
        ]
        docx_bytes = _make_docx_bytes(tables=[table_data])
        result = parse_docx(docx_bytes, filename="test.docx")

        assert isinstance(result, ParseResult)
        assert "1.01" in result.text
        assert len(result.tables) > 0
        # Tables should be in markdown format with pipes
        assert "|" in result.tables[0]


class TestDocxFormattingContext:
    """Test heading/formatting preservation."""

    def test_docx_formatting_context(self):
        """Parse DOCX with headings -> text preserves heading markers."""
        from v2.parsers.docx_parser import parse_docx

        docx_bytes = _make_docx_bytes(
            headings=[
                (1, "Tuerspezifikation"),
                (2, "Brandschutztueren"),
            ],
            paragraphs=["Details zu den Tueren"],
        )
        result = parse_docx(docx_bytes, filename="test.docx")

        assert isinstance(result, ParseResult)
        # Should contain heading markers (# or ##)
        assert "#" in result.text
        assert "Tuerspezifikation" in result.text
        assert "Brandschutztueren" in result.text


class TestDocxErrorHandling:
    """Test error handling for corrupt DOCX."""

    def test_docx_corrupt_file(self):
        """Parse garbage bytes -> returns ParseResult with warning, does not raise."""
        from v2.parsers.docx_parser import parse_docx

        result = parse_docx(b"this is not a docx file", filename="corrupt.docx")

        assert isinstance(result, ParseResult)
        assert len(result.warnings) > 0
        assert result.format == "docx"
