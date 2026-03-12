"""
═══════════════════════════════════════════════════════════════════════════════
Document Parser Service – Production-Grade
Konvertiert Dateien (PDF, Excel, Word, TXT) zu Text mit Error-Handling
═══════════════════════════════════════════════════════════════════════════════
"""

import io
import logging
from pathlib import Path
from typing import Optional, Tuple

from services.error_handler import FileError, ProcessingError, ErrorCode

logger = logging.getLogger(__name__)


class DocumentParser:
    """Hauptklasse für Dokumentenanalyse"""
    
    SUPPORTED_FORMATS = {
        ".pdf": "PDF",
        ".xlsx": "Excel",
        ".xls": "Excel",
        ".xlsm": "Excel",
        ".docx": "Word",
        ".doc": "Word",
        ".txt": "Text"
    }
    
    MAX_TEXT_LENGTH = 100000  # 100k Zeichen max
    
    @staticmethod
    def get_format(filename: str) -> Optional[str]:
        """Gibt Format zurück oder None"""
        ext = Path(filename).suffix.lower()
        return DocumentParser.SUPPORTED_FORMATS.get(ext)
    
    @staticmethod
    def is_supported(filename: str) -> bool:
        """Prüft ob Datei unterstützt wird"""
        return DocumentParser.get_format(filename) is not None


def parse_document_bytes(content: bytes, ext: str) -> str:
    """
    Parst Dokument aus Bytes und gibt Text zurück.
    
    Args:
        content: Datei-Bytes
        ext: Dateiendung (z.B. ".pdf")
        
    Returns:
        Extrahierter Text
        
    Raises:
        FileError: Wenn Format nicht unterstützt oder Parsing fehlschlägt
    """
    if not content:
        raise FileError(
            ErrorCode.FILE_PARSE_ERROR,
            "Datei ist leer",
            filename=ext
        )
    
    ext = ext.lower()
    
    try:
        if ext == ".pdf":
            return _parse_pdf_bytes(content)
        elif ext in (".xlsx", ".xls", ".xlsm"):
            return _parse_excel_bytes(content)
        elif ext in (".docx", ".doc"):
            return _parse_word_bytes(content)
        elif ext == ".txt":
            text = content.decode("utf-8", errors="replace").strip()
            if not text:
                raise FileError(
                    ErrorCode.FILE_PARSE_ERROR,
                    "Textdatei ist leer",
                    filename=ext
                )
            return text
        else:
            raise FileError(
                ErrorCode.INVALID_FILE,
                f"Format '{ext}' wird nicht unterstützt",
                filename=ext
            )
    except FileError:
        raise
    except Exception as e:
        logger.exception(f"Parsing fehlgeschlagen für {ext}")
        raise FileError(
            ErrorCode.FILE_PARSE_ERROR,
            f"Konnte {ext} nicht parsen: {str(e)}",
            filename=ext
        )


def parse_document(file_path: str) -> str:
    """
    Parst Dokumentdatei und gibt Text zurück.
    Wrapper für Rückwärtskompatibilität.
    """
    path = Path(file_path)
    if not path.exists():
        raise FileError(
            ErrorCode.FILE_NOT_FOUND,
            f"Datei nicht gefunden: {file_path}",
            filename=str(path)
        )
    
    try:
        with open(path, "rb") as f:
            content = f.read()
        return parse_document_bytes(content, path.suffix)
    except FileError:
        raise
    except Exception as e:
        logger.exception(f"Datei-Read fehlgeschlagen: {file_path}")
        raise FileError(
            ErrorCode.FILE_PARSE_ERROR,
            f"Konnte Datei nicht lesen: {str(e)}",
            filename=str(path)
        )


def parse_pdf_specs_bytes(content: bytes, max_chars: int = 8000) -> str:
    """
    Parst Spezifikations-PDF mit limitierter Länge.

    Args:
        content: PDF-Bytes
        max_chars: Max Zeichen zum Extrahieren (0 = unbegrenzt)

    Returns:
        Extrahierter Text (limitiert)
    """
    if not content:
        return ""

    # max_chars=0 means unlimited
    effective_limit = max_chars if max_chars > 0 else 999_999_999

    try:
        text_parts = []
        total_chars = 0

        # ── PyMuPDF text extraction (no page cap) ──
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(stream=content, filetype="pdf")
            total_pages = len(doc)
            logger.info(f"[PDF] Spec parsing: {total_pages} pages (PyMuPDF), max_chars={max_chars}")

            for page_num in range(total_pages):
                if total_chars >= effective_limit:
                    break

                try:
                    page = doc[page_num]
                    page_text = page.get_text()
                    if page_text and page_text.strip():
                        text_parts.append(f"--- Seite {page_num + 1} ---\n{page_text}")
                        total_chars += len(page_text)
                except Exception as e:
                    logger.warning(f"Text-Extraktion Seite {page_num + 1} fehlgeschlagen: {e}")

                # Progress logging every 50 pages and on last page
                if (page_num + 1) % 50 == 0 or page_num + 1 == total_pages:
                    logger.info(f"[PDF] Spec text extraction: {page_num + 1}/{total_pages} pages")

            doc.close()

        except ImportError:
            import pdfplumber
            logger.warning("[PDF] PyMuPDF not installed, falling back to pdfplumber for spec parsing")
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                total_pages = len(pdf.pages)
                logger.info(f"[PDF] Spec parsing: {total_pages} pages (pdfplumber fallback), max_chars={max_chars}")
                for page_num, page in enumerate(pdf.pages, 1):
                    if total_chars >= effective_limit:
                        break
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(f"--- Seite {page_num} ---\n{page_text}")
                            total_chars += len(page_text)
                    except Exception as e:
                        logger.warning(f"Text-Extraktion Seite {page_num} fehlgeschlagen: {e}")

        # ── pdfplumber table extraction ──
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    if total_chars >= effective_limit:
                        break
                    try:
                        tables = page.extract_tables()
                        for table in tables or []:
                            if table and total_chars < effective_limit:
                                table_text = _table_to_text(table)
                                if table_text:
                                    text_parts.append(f"[Tabelle Seite {page_num}]\n{table_text}")
                                    total_chars += len(table_text)
                    except Exception as e:
                        logger.debug(f"Tabellen-Extraktion Seite {page_num} fehlgeschlagen: {e}")
        except Exception as e:
            logger.debug(f"[PDF] Table extraction pass failed: {e}")

        text = "\n\n".join(text_parts)
        if max_chars > 0 and len(text) > max_chars:
            text = text[:max_chars] + "\n[... Text gekürzt]"

        logger.info(f"PDF spec parsed: {len(text_parts)} parts, {total_chars} chars extracted")
        return text if text.strip() else "[PDF konnte nicht vollständig gelesen werden]"

    except Exception as e:
        logger.exception(f"PDF-Specs Parsing fehlgeschlagen")
        return f"[PDF konnte nicht gelesen werden: {str(e)}]"


def parse_pdf_specs(file_path: str, max_chars: int = 8000) -> str:
    """Wrapper für Rückwärtskompatibilität"""
    path = Path(file_path)
    if not path.exists():
        return f"[Datei nicht gefunden: {file_path}]"
    
    try:
        with open(path, "rb") as f:
            content = f.read()
        return parse_pdf_specs_bytes(content, max_chars)
    except Exception as e:
        logger.exception(f"PDF-Specs Read fehlgeschlagen: {file_path}")
        return f"[Datei konnte nicht gelesen werden: {str(e)}]"


def parse_pdf_with_vision(content: bytes) -> dict:
    """
    Parse PDF with hybrid analysis (pdfplumber + Vision).
    Returns {"text": str, "positions": list, "method": str, "stats": dict}.
    For use by analyze router when structured extraction is needed.
    """
    try:
        from services.vision_parser import analyze_pdf_hybrid
        return analyze_pdf_hybrid(content)
    except ImportError:
        logger.debug("[PDF] vision_parser not available")
        return {
            "text": _parse_pdf_bytes(content),
            "positions": [],
            "method": "text_only",
            "stats": {},
        }


# ─────────────────────────────────────────────────────────────────────────────
# INTERNE PARSER
# ─────────────────────────────────────────────────────────────────────────────

def _configure_tesseract():
    """Configure Tesseract binary and tessdata paths."""
    import pytesseract
    import os

    # Set Tesseract binary path (Windows default)
    tesseract_cmd = os.environ.get("TESSERACT_CMD")
    if not tesseract_cmd:
        default = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.isfile(default):
            tesseract_cmd = default
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    # Set tessdata path – prefer project-local data/tessdata, then env var
    if not os.environ.get("TESSDATA_PREFIX"):
        local_tessdata = os.path.join(
            os.path.dirname(__file__), "..", "..", "data", "tessdata"
        )
        if os.path.isdir(local_tessdata):
            os.environ["TESSDATA_PREFIX"] = os.path.abspath(local_tessdata)


def _ocr_pdf_bytes(content: bytes, max_pages: int = 30) -> str:
    """
    OCR fallback for scanned PDFs using Tesseract.
    Converts PDF pages to images, then runs OCR.
    Returns extracted text or empty string if OCR is not available.
    """
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        logger.debug("[OCR] pytesseract or Pillow not available")
        return ""

    _configure_tesseract()

    try:
        # Try pdf2image first (best quality)
        try:
            from pdf2image import convert_from_bytes
            images = convert_from_bytes(content, first_page=1, last_page=max_pages, dpi=200)
        except ImportError:
            # Fallback: use pdfplumber to render pages as images
            import pdfplumber
            images = []
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages[:max_pages]:
                    try:
                        img = page.to_image(resolution=200)
                        images.append(img.original)
                    except Exception:
                        continue

        if not images:
            return ""

        # Determine available languages
        ocr_lang = "deu+eng"
        try:
            available = pytesseract.get_languages()
            if "deu" not in available:
                ocr_lang = "eng"
                logger.debug("[OCR] German language not available, using English only")
        except Exception:
            ocr_lang = "eng"

        text_parts = []
        for i, img in enumerate(images, 1):
            try:
                page_text = pytesseract.image_to_string(img, lang=ocr_lang)
                if page_text and page_text.strip():
                    text_parts.append(f"--- Seite {i} (OCR) ---\n{page_text}")
            except Exception as e:
                logger.debug(f"[OCR] Page {i} failed: {e}")
                continue

        return "\n\n".join(text_parts)
    except Exception as e:
        logger.warning(f"[OCR] PDF OCR failed: {e}")
        return ""


def _parse_pdf_bytes(content: bytes) -> str:
    """Extrahiert Text aus PDF mit PyMuPDF (schnell), pdfplumber für Tabellen, OCR-Fallback für gescannte Dokumente"""
    import pdfplumber

    text_parts = []

    try:
        # ── PyMuPDF text extraction (10-50x faster than pdfplumber) ──
        try:
            import fitz  # PyMuPDF

            doc = fitz.open(stream=content, filetype="pdf")
            total_pages = len(doc)

            if total_pages == 0:
                doc.close()
                raise FileError(
                    ErrorCode.FILE_PARSE_ERROR,
                    "PDF hat keine Seiten",
                    filename=".pdf"
                )

            logger.info(f"[PDF] Starting PyMuPDF text extraction: {total_pages} pages")

            for page_num in range(total_pages):
                try:
                    page = doc[page_num]
                    page_text = page.get_text()
                    if page_text and page_text.strip():
                        text_parts.append(f"--- Seite {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    logger.warning(f"Seite {page_num + 1} Text-Extraktion fehlgeschlagen: {e}")
                    continue

                # Progress logging every 50 pages and on last page
                if (page_num + 1) % 50 == 0 or page_num + 1 == total_pages:
                    logger.info(f"[PDF] Text extraction: {page_num + 1}/{total_pages} pages")

            doc.close()

        except ImportError:
            logger.warning("[PDF] PyMuPDF not installed, falling back to pdfplumber")
            # Fallback: pdfplumber text extraction
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                if not pdf.pages:
                    raise FileError(
                        ErrorCode.FILE_PARSE_ERROR,
                        "PDF hat keine Seiten",
                        filename=".pdf"
                    )
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            text_parts.append(f"--- Seite {page_num} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Seite {page_num} Extraktion fehlgeschlagen: {e}")
                        continue

        # ── pdfplumber table extraction (PyMuPDF not good at tables) ──
        try:
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        tables = page.extract_tables()
                        for table in tables or []:
                            if table:
                                table_text = _table_to_text(table)
                                if table_text:
                                    text_parts.append(f"[Tabelle Seite {page_num}]\n{table_text}")
                    except Exception as e:
                        logger.debug(f"Tabellen-Extraktion Seite {page_num} fehlgeschlagen: {e}")
                        continue
        except Exception as e:
            logger.debug(f"[PDF] Table extraction pass failed: {e}")

        result = "\n\n".join(text_parts)

        # OCR fallback for scanned PDFs (no text extracted)
        if not result.strip():
            ocr_text = _ocr_pdf_bytes(content)
            if ocr_text and ocr_text.strip():
                logger.info("[PDF] Using OCR fallback for scanned document")
                return ocr_text[:DocumentParser.MAX_TEXT_LENGTH]
            # Vision fallback for scanned/drawing PDFs
            try:
                from services.vision_parser import analyze_pdf_hybrid
                vision_result = analyze_pdf_hybrid(content)
                if vision_result.get("positions"):
                    pos_texts = []
                    for p in vision_result["positions"]:
                        pos_texts.append(
                            f"Position {p.get('position', p.get('tuer_nr', '?'))}: "
                            f"{p.get('beschreibung', '')} "
                            f"Menge: {p.get('menge', 1)}"
                        )
                    vision_text = "\n".join(pos_texts)
                    if vision_text:
                        logger.info(f"[PDF] Vision extracted {len(vision_result['positions'])} positions")
                        return vision_text[:DocumentParser.MAX_TEXT_LENGTH]
                elif vision_result.get("text"):
                    return vision_result["text"][:DocumentParser.MAX_TEXT_LENGTH]
            except Exception as e:
                logger.debug(f"[PDF] Vision fallback failed: {e}")
            raise FileError(
                ErrorCode.FILE_PARSE_ERROR,
                "PDF ist leer oder konnte nicht gelesen werden (auch OCR fehlgeschlagen)",
                filename=".pdf"
            )

        # If text is suspiciously short, try OCR for additional content
        if len(result.strip()) < 200:
            ocr_text = _ocr_pdf_bytes(content)
            if ocr_text and len(ocr_text.strip()) > len(result.strip()):
                logger.info("[PDF] OCR produced better results than text extraction")
                result = ocr_text
            # If still short, try Vision
            if len(result.strip()) < 200:
                try:
                    from services.vision_parser import analyze_pdf_hybrid
                    vision_result = analyze_pdf_hybrid(content)
                    if vision_result.get("positions"):
                        pos_texts = []
                        for p in vision_result["positions"]:
                            pos_texts.append(
                                f"Position {p.get('position', p.get('tuer_nr', '?'))}: "
                                f"{p.get('beschreibung', '')} "
                                f"Menge: {p.get('menge', 1)}"
                            )
                        vision_text = "\n".join(pos_texts)
                        if vision_text and len(vision_text) > len(result.strip()):
                            logger.info(f"[PDF] Vision extracted {len(vision_result['positions'])} positions (short text fallback)")
                            result = vision_text
                    elif vision_result.get("text") and len(vision_result["text"].strip()) > len(result.strip()):
                        result = vision_result["text"]
                except Exception as e:
                    logger.debug(f"[PDF] Vision fallback failed: {e}")

        return result[:DocumentParser.MAX_TEXT_LENGTH]

    except FileError:
        raise
    except Exception as e:
        raise FileError(
            ErrorCode.FILE_PARSE_ERROR,
            f"PDF-Parsing fehlgeschlagen: {str(e)}",
            filename=".pdf"
        )


def _parse_excel_bytes(content: bytes) -> str:
    """Extrahiert Text aus Excel mit pandas"""
    import pandas as pd
    
    text_parts = []
    
    try:
        buf = io.BytesIO(content)
        xl = pd.ExcelFile(buf)
        try:
            if not xl.sheet_names:
                raise FileError(
                    ErrorCode.FILE_PARSE_ERROR,
                    "Excel hat keine Tabellenblätter",
                    filename=".xlsx"
                )

            for sheet_name in xl.sheet_names:
                try:
                    df = pd.read_excel(xl, sheet_name=sheet_name, dtype=str)
                    df = df.fillna("")

                    if df.empty:
                        continue

                    text_parts.append(f"=== {sheet_name} ===")
                    text_parts.append(df.to_string(index=False))

                except Exception as e:
                    logger.warning(f"Sheet '{sheet_name}' Parsing fehlgeschlagen: {e}")
                    continue
        finally:
            xl.close()

        result = "\n\n".join(text_parts)
        if not result.strip():
            raise FileError(
                ErrorCode.FILE_PARSE_ERROR,
                "Excel ist leer",
                filename=".xlsx"
            )

        return result[:DocumentParser.MAX_TEXT_LENGTH]

    except FileError:
        raise
    except Exception as e:
        raise FileError(
            ErrorCode.FILE_PARSE_ERROR,
            f"Excel-Parsing fehlgeschlagen: {str(e)}",
            filename=".xlsx"
        )


def _parse_word_bytes(content: bytes) -> str:
    """Extrahiert Text aus Word mit python-docx"""
    from docx import Document
    
    text_parts = []
    
    try:
        doc = Document(io.BytesIO(content))
        
        # Paragraphen
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Tabellen
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text_parts.append(" | ".join(cells))
        
        result = "\n".join(text_parts)
        if not result.strip():
            raise FileError(
                ErrorCode.FILE_PARSE_ERROR,
                "Word-Dokument ist leer",
                filename=".docx"
            )
        
        return result[:DocumentParser.MAX_TEXT_LENGTH]
    
    except FileError:
        raise
    except Exception as e:
        raise FileError(
            ErrorCode.FILE_PARSE_ERROR,
            f"Word-Parsing fehlgeschlagen: {str(e)}",
            filename=".docx"
        )


def _table_to_text(table: list) -> str:
    """Konvertiert Tabelle (Liste von Listen) zu Text"""
    if not table:
        return ""
    
    rows = []
    for row in table:
        if row:
            cells = [str(cell).strip() if cell else "" for cell in row]
            rows.append(" | ".join(cells))
    
    return "\n".join(rows)


if __name__ == "__main__":
    # Test
    import sys
    if len(sys.argv) > 1:
        text = parse_document(sys.argv[1])
        print(f"Extrahiert: {len(text)} Zeichen")
        print(text[:500])
