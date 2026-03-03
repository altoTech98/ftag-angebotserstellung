"""
Offer Generator – Creates Excel and Word offers and gap reports in memory.
All functions return bytes (no disk writes).
"""

import io
from datetime import datetime, timedelta


def generate_offer_excel(
    offer_text: str,
    matched_positions: list,
    requirements: dict,
    offer_id: str,
) -> bytes:
    """Generate an Excel offer in memory. Returns raw bytes."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Angebot"

    # Colors
    header_fill = PatternFill("solid", fgColor="1F3A5F")
    subheader_fill = PatternFill("solid", fgColor="2E6DA4")
    alt_row_fill = PatternFill("solid", fgColor="EBF3FB")
    total_fill = PatternFill("solid", fgColor="F0F0F0")

    # Fonts
    title_font = Font(name="Arial", size=16, bold=True, color="1F3A5F")
    header_font = Font(name="Arial", size=10, bold=True, color="FFFFFF")
    normal_font = Font(name="Arial", size=10)
    bold_font = Font(name="Arial", size=10, bold=True)
    small_font = Font(name="Arial", size=9, color="666666")

    thin_border = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )

    # Column widths
    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 40
    ws.column_dimensions["C"].width = 10
    ws.column_dimensions["D"].width = 8
    ws.column_dimensions["E"].width = 16
    ws.column_dimensions["F"].width = 16

    row = 1

    # Header: Company
    ws.merge_cells(f"A{row}:F{row}")
    ws[f"A{row}"] = "Frank Türen AG"
    ws[f"A{row}"].font = title_font
    ws[f"A{row}"].alignment = Alignment(horizontal="left")
    row += 1

    ws.merge_cells(f"A{row}:F{row}")
    ws[f"A{row}"] = "Industriestrasse 12 · 6374 Buochs NW · Tel. 041 620 76 76 · www.franktueren.ch"
    ws[f"A{row}"].font = small_font
    row += 1

    ws.row_dimensions[row].height = 10
    row += 1

    # Offer info
    today = datetime.now()
    valid_until = today + timedelta(days=90)
    offer_number = f"AG-{today.strftime('%Y%m')}-{offer_id[:6].upper()}"

    project_name = requirements.get("projekt", "Ausschreibung")
    client_name = requirements.get("auftraggeber", "")

    info_data = [
        ("Angebotsnummer:", offer_number),
        ("Datum:", today.strftime("%d.%m.%Y")),
        ("Gültig bis:", valid_until.strftime("%d.%m.%Y")),
        ("Projekt:", project_name),
    ]
    if client_name:
        info_data.append(("Auftraggeber:", client_name))

    for label, value in info_data:
        ws[f"A{row}"] = label
        ws[f"A{row}"].font = bold_font
        ws.merge_cells(f"B{row}:F{row}")
        ws[f"B{row}"] = value
        ws[f"B{row}"].font = normal_font
        row += 1

    row += 1

    # Greeting
    ws.merge_cells(f"A{row}:F{row}")
    ws[f"A{row}"] = "Sehr geehrte Damen und Herren"
    ws[f"A{row}"].font = normal_font
    row += 1

    ws.merge_cells(f"A{row}:F{row}")
    ws[f"A{row}"] = (
        f"Wir erlauben uns, Ihnen für das Projekt «{project_name}» folgendes Angebot zu unterbreiten:"
    )
    ws[f"A{row}"].font = normal_font
    ws[f"A{row}"].alignment = Alignment(wrap_text=True)
    ws.row_dimensions[row].height = 20
    row += 1

    row += 1

    # Table header
    headers = ["Pos.", "Beschreibung", "Menge", "Einheit", "Einzelpreis CHF", "Gesamtpreis CHF"]
    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=row, column=col_idx, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = thin_border
    ws.row_dimensions[row].height = 25
    row += 1

    # Positions
    total_sum = 0

    for i, pos_data in enumerate(matched_positions):
        pos = pos_data.get("original_position", pos_data)
        position_nr = pos.get("position", str(i + 1))
        menge = pos.get("menge", 1) or 1
        einheit = pos.get("einheit", "Stk")

        tuertyp = pos.get("tuertyp", "Tür")
        breite = pos.get("breite")
        hoehe = pos.get("hoehe")
        brandschutz = pos.get("brandschutz")
        einbruchschutz = pos.get("einbruchschutz")
        beschreibung = pos.get("beschreibung", "")

        desc_parts = [beschreibung or tuertyp or "Tür"]
        if breite and hoehe:
            desc_parts.append(f"B {breite} × H {hoehe} mm")
        if brandschutz:
            desc_parts.append(f"Brandschutz {brandschutz}")
        if einbruchschutz:
            desc_parts.append(f"Einbruchschutz {einbruchschutz}")

        description = " | ".join(filter(None, desc_parts))

        unit_price = _estimate_price(pos)
        total_price = unit_price * int(menge)
        total_sum += total_price

        fill = alt_row_fill if i % 2 == 0 else PatternFill("solid", fgColor="FFFFFF")

        row_data = [position_nr, description, menge, einheit, unit_price, total_price]
        alignments = ["center", "left", "center", "center", "right", "right"]

        for col_idx, (value, align) in enumerate(zip(row_data, alignments), 1):
            cell = ws.cell(row=row, column=col_idx, value=value)
            cell.font = normal_font
            cell.fill = fill
            cell.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
            cell.border = thin_border
            if col_idx in (5, 6) and isinstance(value, (int, float)):
                cell.number_format = "#,##0.00"

        ws.row_dimensions[row].height = 30
        row += 1

    # Totals
    row += 1
    mwst = total_sum * 0.081
    grand_total = total_sum + mwst

    for label, value in [
        ("Zwischensumme CHF", total_sum),
        ("MwSt 8.1% CHF", mwst),
        ("", None),
    ]:
        ws.merge_cells(f"A{row}:E{row}")
        ws[f"A{row}"] = label
        ws[f"A{row}"].font = normal_font
        ws[f"A{row}"].alignment = Alignment(horizontal="right")
        ws[f"A{row}"].fill = total_fill
        if value is not None:
            ws[f"F{row}"] = value
            ws[f"F{row}"].font = normal_font
            ws[f"F{row}"].number_format = "#,##0.00"
            ws[f"F{row}"].alignment = Alignment(horizontal="right")
            ws[f"F{row}"].fill = total_fill
            ws[f"F{row}"].border = thin_border
        row += 1

    # Grand total
    ws.merge_cells(f"A{row}:E{row}")
    ws[f"A{row}"] = "GESAMTBETRAG CHF (inkl. MwSt)"
    ws[f"A{row}"].font = Font(name="Arial", size=11, bold=True)
    ws[f"A{row}"].fill = subheader_fill
    ws[f"A{row}"].alignment = Alignment(horizontal="right")
    ws[f"F{row}"] = grand_total
    ws[f"F{row}"].font = Font(name="Arial", size=11, bold=True, color="FFFFFF")
    ws[f"F{row}"].fill = subheader_fill
    ws[f"F{row}"].number_format = "#,##0.00"
    ws[f"F{row}"].alignment = Alignment(horizontal="right")
    ws[f"F{row}"].border = thin_border
    row += 2

    # Conditions
    conditions = [
        "Zahlungsbedingungen: 30 Tage netto ab Rechnungsdatum",
        "Lieferzeit: nach Vereinbarung, ca. 6–8 Wochen ab Auftragserteilung",
        "Preise: Richtpreise exkl. Montage, CHF exkl. MwSt",
        "Gültigkeit: 90 Tage ab Angebotsdatum",
    ]
    for cond in conditions:
        ws.merge_cells(f"A{row}:F{row}")
        ws[f"A{row}"] = cond
        ws[f"A{row}"].font = small_font
        row += 1

    row += 1
    ws.merge_cells(f"A{row}:F{row}")
    ws[f"A{row}"] = "Wir danken Ihnen für Ihr Interesse und freuen uns auf Ihre Auftragserteilung."
    ws[f"A{row}"].font = normal_font
    row += 2

    ws.merge_cells(f"A{row}:F{row}")
    ws[f"A{row}"] = "Freundliche Grüsse"
    ws[f"A{row}"].font = normal_font
    row += 2

    ws.merge_cells(f"A{row}:F{row}")
    ws[f"A{row}"] = "Frank Türen AG"
    ws[f"A{row}"].font = bold_font
    row += 1

    ws.merge_cells(f"A{row}:F{row}")
    ws[f"A{row}"] = "Buochs NW"
    ws[f"A{row}"].font = normal_font

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_gap_report_excel(
    gap_report_text: str,
    unmatched_positions: list,
    partial_positions: list,
    requirements: dict,
    report_id: str,
) -> bytes:
    """Generate a gap report Excel in memory. Returns raw bytes."""
    import openpyxl
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Gap Report"

    # Styles
    title_font = Font(name="Arial", size=16, bold=True, color="C0392B")
    header_font = Font(name="Arial", size=10, bold=True, color="FFFFFF")
    normal_font = Font(name="Arial", size=10)
    bold_font = Font(name="Arial", size=10, bold=True)
    small_font = Font(name="Arial", size=9, color="666666")

    red_fill = PatternFill("solid", fgColor="C0392B")
    orange_fill = PatternFill("solid", fgColor="E67E22")
    light_red_fill = PatternFill("solid", fgColor="FADBD8")
    light_orange_fill = PatternFill("solid", fgColor="FDEBD0")

    thin_border = Border(
        left=Side(style="thin", color="CCCCCC"),
        right=Side(style="thin", color="CCCCCC"),
        top=Side(style="thin", color="CCCCCC"),
        bottom=Side(style="thin", color="CCCCCC"),
    )

    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 35
    ws.column_dimensions["C"].width = 15
    ws.column_dimensions["D"].width = 15
    ws.column_dimensions["E"].width = 35

    row = 1

    ws.merge_cells(f"A{row}:E{row}")
    ws[f"A{row}"] = "GAP-REPORT – Frank Türen AG"
    ws[f"A{row}"].font = title_font
    row += 1

    ws.merge_cells(f"A{row}:E{row}")
    ws[f"A{row}"] = f"Projekt: {requirements.get('projekt', 'Ausschreibung')} · Erstellt: {datetime.now().strftime('%d.%m.%Y')}"
    ws[f"A{row}"].font = small_font
    row += 2

    # Summary
    total = len(requirements.get("positionen", []))
    n_unmatched = len(unmatched_positions)
    n_partial = len(partial_positions)

    summary_data = [
        ("Gesamte Positionen:", str(total)),
        ("Nicht erfüllbar:", f"{n_unmatched} Positionen"),
        ("Teilweise erfüllbar:", f"{n_partial} Positionen"),
        ("Empfehlung:", "Sonderanfertigung oder Alternativprodukt prüfen"),
    ]
    for label, value in summary_data:
        ws[f"A{row}"] = label
        ws[f"A{row}"].font = bold_font
        ws.merge_cells(f"B{row}:E{row}")
        ws[f"B{row}"] = value
        ws[f"B{row}"].font = normal_font
        row += 1

    row += 1

    # Unmatched positions
    if unmatched_positions:
        ws.merge_cells(f"A{row}:E{row}")
        ws[f"A{row}"] = "NICHT ERFÜLLBARE POSITIONEN"
        ws[f"A{row}"].font = header_font
        ws[f"A{row}"].fill = red_fill
        ws[f"A{row}"].alignment = Alignment(horizontal="center")
        row += 1

        headers = ["Pos.", "Beschreibung", "Anforderung", "Status", "Empfehlung"]
        for col_idx, h in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col_idx, value=h)
            cell.font = header_font
            cell.fill = red_fill
            cell.alignment = Alignment(horizontal="center")
            cell.border = thin_border
        row += 1

        for item in unmatched_positions:
            pos = item.get("original_position", item)
            req_parts = []
            if pos.get("brandschutz"):
                req_parts.append(f"Brandschutz: {pos['brandschutz']}")
            if pos.get("einbruchschutz"):
                req_parts.append(f"Einbruch: {pos['einbruchschutz']}")
            if pos.get("tuertyp"):
                req_parts.append(f"Typ: {pos['tuertyp']}")

            row_data = [
                item.get("position", "?"),
                item.get("beschreibung", pos.get("beschreibung", "")),
                "\n".join(req_parts) or "Spezifische Anforderungen",
                "Nicht verfügbar",
                item.get("reason", "Sonderanfertigung prüfen"),
            ]
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row, column=col_idx, value=value)
                cell.font = normal_font
                cell.fill = light_red_fill
                cell.alignment = Alignment(vertical="center", wrap_text=True)
                cell.border = thin_border
            ws.row_dimensions[row].height = 35
            row += 1

    row += 1

    # Partial positions
    if partial_positions:
        ws.merge_cells(f"A{row}:E{row}")
        ws[f"A{row}"] = "TEILWEISE ERFÜLLBARE POSITIONEN"
        ws[f"A{row}"].font = header_font
        ws[f"A{row}"].fill = orange_fill
        ws[f"A{row}"].alignment = Alignment(horizontal="center")
        row += 1

        headers = ["Pos.", "Beschreibung", "Anforderung", "Status", "Hinweis"]
        for col_idx, h in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col_idx, value=h)
            cell.font = header_font
            cell.fill = orange_fill
            cell.alignment = Alignment(horizontal="center")
            cell.border = thin_border
        row += 1

        for item in partial_positions:
            pos = item.get("original_position", item)
            req_parts = []
            if pos.get("brandschutz"):
                req_parts.append(f"Brandschutz: {pos['brandschutz']}")
            if pos.get("einbruchschutz"):
                req_parts.append(f"Einbruch: {pos['einbruchschutz']}")

            row_data = [
                item.get("position", "?"),
                item.get("beschreibung", pos.get("beschreibung", "")),
                "\n".join(req_parts) or "Teilweise spezifiziert",
                "Rückfrage nötig",
                item.get("reason", "Technische Rückfrage empfohlen"),
            ]
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row, column=col_idx, value=value)
                cell.font = normal_font
                cell.fill = light_orange_fill
                cell.alignment = Alignment(vertical="center", wrap_text=True)
                cell.border = thin_border
            ws.row_dimensions[row].height = 35
            row += 1

    row += 2

    # Recommendations
    ws.merge_cells(f"A{row}:E{row}")
    ws[f"A{row}"] = "HANDLUNGSEMPFEHLUNGEN"
    ws[f"A{row}"].font = bold_font
    row += 1

    recommendations = [
        "1. Nicht erfüllbare Positionen intern prüfen (Sonderanfertigung möglich?)",
        "2. Technische Rückfrage beim Auftraggeber für Alternativspezifikationen",
        "3. Lieferantenpartner für spezielle Anforderungen kontaktieren",
        "4. Angebot für erfüllbare Positionen separat einreichen (Teilangebot)",
    ]
    for rec in recommendations:
        ws.merge_cells(f"A{row}:E{row}")
        ws[f"A{row}"] = rec
        ws[f"A{row}"].font = normal_font
        row += 1

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_offer_word(
    offer_text: str,
    matched_positions: list,
    requirements: dict,
    offer_id: str,
) -> bytes:
    """Generate a Word (.docx) offer in memory. Returns raw bytes."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    today = datetime.now()
    valid_until = today + timedelta(days=90)
    offer_number = f"AG-{today.strftime('%Y%m')}-{offer_id[:6].upper()}"
    project_name = requirements.get("projekt", "Ausschreibung")
    client_name = requirements.get("auftraggeber", "")

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    # Header
    h = doc.add_heading("Frank Türen AG", 0)
    h.runs[0].font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sub = doc.add_paragraph("Industriestrasse 12 · 6374 Buochs NW · Tel. 041 620 76 76 · www.franktueren.ch")
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    # Offer meta info table
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = "Table Grid"
    meta = [
        ("Angebotsnummer:", offer_number),
        ("Datum:", today.strftime("%d.%m.%Y")),
        ("Gültig bis:", valid_until.strftime("%d.%m.%Y")),
        ("Projekt:", project_name),
    ]
    if client_name:
        meta.append(("Auftraggeber:", client_name))
    for label, value in meta:
        row = info_table.add_row()
        row.cells[0].text = label
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = value
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        set_cell_bg(row.cells[0], "EBF3FB")
    info_table.columns[0].width = Cm(4)
    info_table.columns[1].width = Cm(12)

    doc.add_paragraph()
    greeting = doc.add_paragraph("Sehr geehrte Damen und Herren")
    intro = doc.add_paragraph(
        f"Wir erlauben uns, Ihnen für das Projekt «{project_name}» folgendes Angebot zu unterbreiten:"
    )
    for p in (greeting, intro):
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # Positions table
    pos_table = doc.add_table(rows=1, cols=6)
    pos_table.style = "Table Grid"
    headers = ["Pos.", "Beschreibung", "Menge", "Einheit", "Einzelpreis CHF", "Gesamtpreis CHF"]
    hdr_row = pos_table.rows[0]
    for i, (cell, hdr) in enumerate(zip(hdr_row.cells, headers)):
        cell.text = hdr
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "1F3A5F")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    total_sum = 0.0
    for i, pos_data in enumerate(matched_positions):
        pos = pos_data.get("original_position", pos_data)
        menge = int(pos.get("menge") or 1)
        einheit = pos.get("einheit", "Stk")
        tuertyp = pos.get("tuertyp", "Tür")
        breite = pos.get("breite")
        hoehe = pos.get("hoehe")
        brandschutz = pos.get("brandschutz")
        einbruchschutz = pos.get("einbruchschutz")
        beschreibung = pos.get("beschreibung", "")

        desc_parts = [beschreibung or tuertyp]
        if breite and hoehe:
            desc_parts.append(f"B {breite} × H {hoehe} mm")
        if brandschutz:
            desc_parts.append(f"Brandschutz {brandschutz}")
        if einbruchschutz:
            desc_parts.append(f"Einbruchschutz {einbruchschutz}")

        unit_price = _estimate_price(pos)
        total_price = unit_price * menge
        total_sum += total_price
        fill = "EBF3FB" if i % 2 == 0 else "FFFFFF"

        row = pos_table.add_row()
        values = [pos.get("position", str(i+1)), " | ".join(desc_parts), str(menge), einheit,
                  f"{unit_price:,.2f}", f"{total_price:,.2f}"]
        aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT]
        for cell, val, align in zip(row.cells, values, aligns):
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = align
            set_cell_bg(cell, fill)

    # Totals
    mwst = total_sum * 0.081
    grand_total = total_sum + mwst
    for label, value in [("Zwischensumme CHF", total_sum), ("MwSt 8.1% CHF", mwst), ("GESAMTBETRAG CHF (inkl. MwSt)", grand_total)]:
        row = pos_table.add_row()
        merged = row.cells[0].merge(row.cells[4])
        merged.text = label
        merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        is_total = "GESAMT" in label
        merged.paragraphs[0].runs[0].bold = is_total
        merged.paragraphs[0].runs[0].font.size = Pt(10 if is_total else 9)
        row.cells[5].text = f"{value:,.2f}"
        row.cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        row.cells[5].paragraphs[0].runs[0].bold = is_total
        row.cells[5].paragraphs[0].runs[0].font.size = Pt(10 if is_total else 9)
        if is_total:
            set_cell_bg(merged, "2E6DA4")
            merged.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_bg(row.cells[5], "2E6DA4")
            row.cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()
    for cond in ["Zahlungsbedingungen: 30 Tage netto ab Rechnungsdatum",
                 "Preise: Richtpreise exkl. Montage, CHF exkl. MwSt",
                 "Lieferzeit: ca. 6–8 Wochen ab Auftragserteilung"]:
        p = doc.add_paragraph(cond)
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()
    doc.add_paragraph("Wir danken Ihnen für Ihr Interesse und freuen uns auf Ihre Auftragserteilung.").runs[0].font.size = Pt(10)
    doc.add_paragraph()
    doc.add_paragraph("Freundliche Grüsse").runs[0].font.size = Pt(10)
    doc.add_paragraph()
    p = doc.add_paragraph("Frank Türen AG")
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(10)
    doc.add_paragraph("Buochs NW").runs[0].font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def generate_gap_report_word(
    gap_report_text: str,
    unmatched_positions: list,
    partial_positions: list,
    requirements: dict,
    report_id: str,
) -> bytes:
    """Generate a Word (.docx) gap report in memory. Returns raw bytes."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    # Title
    h = doc.add_heading("GAP-REPORT – Frank Türen AG", 0)
    h.runs[0].font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C)
    sub = doc.add_paragraph(
        f"Projekt: {requirements.get('projekt', 'Ausschreibung')} · Erstellt: {datetime.now().strftime('%d.%m.%Y')}"
    )
    sub.runs[0].font.size = Pt(9)
    sub.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    doc.add_paragraph()

    # Summary
    total = len(requirements.get("positionen", []))
    summary_data = [
        ("Gesamte Positionen:", str(total)),
        ("Nicht erfüllbar:", f"{len(unmatched_positions)} Positionen"),
        ("Teilweise erfüllbar:", f"{len(partial_positions)} Positionen"),
        ("Empfehlung:", "Sonderanfertigung oder Alternativprodukt prüfen"),
    ]
    for label, value in summary_data:
        p = doc.add_paragraph()
        run_l = p.add_run(f"{label} ")
        run_l.bold = True
        run_l.font.size = Pt(10)
        run_v = p.add_run(value)
        run_v.font.size = Pt(10)

    doc.add_paragraph()

    def add_positions_section(title, positions, fill_color, header_color):
        if not positions:
            return
        h = doc.add_heading(title, 2)
        h.runs[0].font.color.rgb = RGBColor(
            int(header_color[:2], 16), int(header_color[2:4], 16), int(header_color[4:], 16)
        )
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for cell, hdr in zip(table.rows[0].cells, ["Pos.", "Beschreibung", "Anforderung", "Status", "Empfehlung"]):
            cell.text = hdr
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            set_cell_bg(cell, header_color)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for item in positions:
            pos = item.get("original_position", item)
            req_parts = []
            if pos.get("brandschutz"):
                req_parts.append(f"Brandschutz: {pos['brandschutz']}")
            if pos.get("einbruchschutz"):
                req_parts.append(f"Einbruch: {pos['einbruchschutz']}")
            if pos.get("tuertyp"):
                req_parts.append(f"Typ: {pos['tuertyp']}")
            row = table.add_row()
            values = [item.get("position", "?"), item.get("beschreibung", ""),
                      "\n".join(req_parts) or "Spezifisch",
                      "Nicht verfügbar" if "unmatched" in str(item.get("status","")) else "Rückfrage",
                      item.get("reason", "Sonderanfertigung prüfen")]
            for cell, val in zip(row.cells, values):
                cell.text = val
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                set_cell_bg(cell, fill_color)
        doc.add_paragraph()

    add_positions_section("Nicht erfüllbare Positionen", unmatched_positions, "FADBD8", "C0392B")
    add_positions_section("Teilweise erfüllbare Positionen", partial_positions, "FDEBD0", "E67E22")

    doc.add_heading("Handlungsempfehlungen", 2)
    for rec in ["Nicht erfüllbare Positionen intern prüfen (Sonderanfertigung möglich?)",
                "Technische Rückfrage beim Auftraggeber für Alternativspezifikationen",
                "Lieferantenpartner für spezielle Anforderungen kontaktieren",
                "Angebot für erfüllbare Positionen separat einreichen (Teilangebot)"]:
        p = doc.add_paragraph(rec, style="List Number")
        p.runs[0].font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _estimate_price(position: dict) -> float:
    """
    Estimate a price for a door position based on its specifications.
    Returns price in CHF (excl. VAT, excl. installation).
    """
    base_price = 950.0

    tuertyp = (position.get("tuertyp") or "").lower()
    brandschutz = (position.get("brandschutz") or "").upper()
    einbruchschutz = (position.get("einbruchschutz") or "").upper()
    breite = position.get("breite") or 900
    hoehe = position.get("hoehe") or 2100

    if "stahl" in tuertyp:
        base_price = 1100.0
    elif "alu" in tuertyp or "aluminium" in tuertyp:
        base_price = 1400.0
    elif "holz" in tuertyp:
        base_price = 850.0

    fire_premiums = {
        "T30": 400, "EI30": 400,
        "T60": 700, "EI60": 700,
        "T90": 1100, "EI90": 1100,
        "T120": 1500, "EI120": 1500,
    }
    for cls, premium in fire_premiums.items():
        if cls in brandschutz:
            base_price += premium
            break

    burglary_premiums = {
        "RC2": 500, "WK2": 500,
        "RC3": 900, "WK3": 900,
        "RC4": 1500, "WK4": 1500,
    }
    for cls, premium in burglary_premiums.items():
        if cls in einbruchschutz:
            base_price += premium
            break

    area = (breite * hoehe) / (900 * 2100)
    if area > 1.3:
        base_price *= 1.2
    elif area > 1.6:
        base_price *= 1.4

    return round(base_price, 2)
