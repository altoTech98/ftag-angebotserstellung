"""
DOCX Parser - Extracts text from Word documents with formatting context.

Preserves heading hierarchy with Markdown markers (# ## ###).
Extracts tables in Markdown format.
Never raises exceptions to callers.
"""

import io
import logging

from v2.parsers.base import ParseResult

logger = logging.getLogger(__name__)


def parse_docx(content: bytes, filename: str = "") -> ParseResult:
    """Parse DOCX bytes into a ParseResult.

    Args:
        content: Raw DOCX file bytes.
        filename: Original filename for provenance tracking.

    Returns:
        ParseResult with extracted text and metadata. Never raises.
    """
    if not content:
        return ParseResult(
            text="",
            format="docx",
            page_count=0,
            warnings=["Empty file provided"],
            metadata={},
            source_file=filename,
        )

    try:
        from docx import Document

        doc = Document(io.BytesIO(content))
        text_parts = []
        tables_md = []
        styles_found = set()

        # Track paragraph positions and table positions for interleaving
        # In python-docx, we iterate body elements to get correct order
        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                # It's a paragraph
                para = _find_paragraph(doc, element)
                if para is not None:
                    style_name = para.style.name if para.style else "Normal"
                    styles_found.add(style_name)
                    text = para.text.strip()
                    if text:
                        prefix = _heading_prefix(style_name)
                        text_parts.append(f"{prefix}{text}" if prefix else text)

            elif tag == "tbl":
                # It's a table
                table = _find_table(doc, element)
                if table is not None:
                    md = _table_to_markdown(table)
                    if md:
                        tables_md.append(md)
                        text_parts.append(md)

        full_text = "\n\n".join(text_parts)

        metadata = {
            "styles_found": sorted(styles_found),
        }

        return ParseResult(
            text=full_text,
            format="docx",
            page_count=1,  # DOCX doesn't have page concept at parse time
            warnings=[],
            metadata=metadata,
            source_file=filename,
            tables=tables_md,
        )

    except Exception as e:
        logger.warning(f"[DOCX] Parsing failed for {filename}: {e}")
        return ParseResult(
            text="",
            format="docx",
            page_count=0,
            warnings=[f"DOCX parsing failed: {str(e)}"],
            metadata={},
            source_file=filename,
        )


def _find_paragraph(doc, element):
    """Find the Paragraph object matching a given XML element."""
    for para in doc.paragraphs:
        if para._element is element:
            return para
    return None


def _find_table(doc, element):
    """Find the Table object matching a given XML element."""
    for table in doc.tables:
        if table._element is element:
            return table
    return None


def _heading_prefix(style_name: str) -> str:
    """Convert a style name to a Markdown heading prefix.

    Examples:
        'Heading 1' -> '# '
        'Heading 2' -> '## '
        'Title' -> '# '
        'Normal' -> ''
    """
    if not style_name:
        return ""

    name_lower = style_name.lower()

    if "heading" in name_lower:
        # Extract heading level number
        for char in style_name:
            if char.isdigit():
                level = int(char)
                return "#" * level + " "
        return "# "  # Default to H1 if no number found

    if name_lower == "title":
        return "# "

    if name_lower == "subtitle":
        return "## "

    return ""


def _table_to_markdown(table) -> str:
    """Convert a python-docx Table to markdown format."""
    if not table.rows:
        return ""

    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        rows.append("| " + " | ".join(cells) + " |")
        if i == 0:
            # Add separator row after header
            rows.append("| " + " | ".join(["---"] * len(cells)) + " |")

    return "\n".join(rows)
